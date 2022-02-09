from django.shortcuts import render
from django.http import HttpResponse
from rest_framework.views import APIView
from django.http import StreamingHttpResponse
import io
import mammoth
from docx import Document
from django.templatetags.static import static
import os
from django.http import FileResponse
import io



THIS_DIR = os.path.dirname(__file__)
filedir = THIS_DIR + '/guidelines.docx'
pdfdir = THIS_DIR + '/schizo.pdf'
supportdir = THIS_DIR + '/support.pdf'
lossdir = THIS_DIR + '/loss.pdf'

def pdf_view(request):
    with open(pdfdir, 'rb') as pdf:
        response = HttpResponse(pdf.read(),content_type='application/pdf')
        response['Content-Disposition'] = 'filename=some_file.pdf'
        return response

def loss_view(request):
    with open(lossdir, 'rb') as pdf:
        response = HttpResponse(pdf.read(),content_type='application/pdf')
        response['Content-Disposition'] = 'filename=some_file.pdf'
        return response

def support_view(request):
    with open(supportdir, 'rb') as pdf:
        response = HttpResponse(pdf.read(),content_type='application/pdf')
        response['Content-Disposition'] = 'filename=some_file.pdf'
        return response




def documents_home(request):
    print(THIS_DIR)
    with open(filedir, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        text = result.value
        text += "<style>p { font-size: 50px; }li { font-size: 30px; }</style>"

    context = {}
    context['docx'] = f"{text}"
    template = 'documents.html'
    return render(request, template, context)


class ExportDocx(APIView):
    def build_document(self):
        document = Document()
        document.add_heading("This is a header")
        document.add_paragraph("This is a normal style paragraph")
        document.add_paragraph(
            "<p>Nice to see Prep note 2</p><ul>\n<li>Prep note 2 content 1</li><li>Prep note 2 content 2</li></ul>")
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.italic = True
        run.add_text("text will have italic style")
        run.add_break()
        return document

    def get(self, request, *args, **kwargs):
        document = self.build_document()
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        response = StreamingHttpResponse(streaming_content=buffer,
                                         content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment;filename=Test.docx'
        response["Content-Encoding"] = 'UTF-8'
        return response
# Create your views here.
